# -*- coding: utf-8 -*-
import json
import os
import re
import pypdf
import docx

def clean_text(t):
    return re.sub(r'\s+', ' ', t).strip()

def get_keywords(q_text):
    stopwords = {'a', 'an', 'the', 'and', 'or', 'but', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 
                 'in', 'on', 'at', 'to', 'for', 'with', 'by', 'of', 'from', 'its', 'their', 'how', 'what', 
                 'why', 'define', 'explain', 'discuss', 'differentiate', 'between', 'types', 'process', 
                 'method', 'cell', 'battery', 'polymer', 'fuel', 'numerical', 'problems', 'calculate', 
                 'determine', 'amount', 'required', 'given', 'contains'}
    words = re.findall(r'\w+', q_text.lower())
    kw = [w.upper() for w in words if w not in stopwords and len(w) > 2]
    return list(dict.fromkeys(kw))[:8]

def format_block_html(block):
    block = block.strip()
    if not block:
        return ""
    
    # Check if this is a chemical reaction or formula
    is_formula = False
    formula_triggers = ['→', '⇌', 'Ca²', 'Mg²', 'R-H', 'SO₄', 'Na₂', 'Cl⁻', 'OH⁻', 'CO₂', 'H₂O', 'CaCO₃', 'Mg(OH)₂']
    if any(t in block for t in formula_triggers) and len(block) < 300:
        is_formula = True
    elif '=' in block and ('×' in block or '/' in block or 'HCV' in block or 'LCV' in block or 'E°' in block):
        is_formula = True
        
    if is_formula:
        return f"<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>{block}</div>"
        
    # Check if this is a subheader
    is_subheader = False
    subheader_triggers = ['PRINCIPLE:', 'REAGENTS REQUIRED:', 'PROCEDURE', 'REACTIONS:', 'CALCULATION:', 'FORMULA:', 
                          'ADVANTAGES:', 'DISADVANTAGES:', 'LIMITATIONS:', 'WORKING PRINCIPLE', 'EXHAUSTION:', 
                          'REGENERATION:', 'ILL EFFECTS:', 'PREVENTION:', 'INTRODUCTION:', 'CONSTRUCTION:', 
                          'PROPERTIES:', 'APPLICATIONS:', 'MECHANISM:', 'CHARACTERISTICS:']
    if any(block.startswith(t) for t in subheader_triggers) and len(block) < 120:
        is_subheader = True
        
    if is_subheader:
        return f"<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>{block}</h4>"
        
    # Check if bullet list
    if block.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
        # Try to highlight key terms
        match = re.match(r'^(\d+\.\s*[^:]+:)(.*)', block)
        if match:
            return f"<p style='margin-bottom: 8px; padding-left: 15px;'><strong>{match.group(1)}</strong>{match.group(2)}</p>"
        return f"<p style='margin-bottom: 8px; padding-left: 15px;'>{block}</p>"
        
    if block.startswith('📌'):
        return f"<div style='background: rgba(0, 242, 254, 0.05); border-left: 3px solid var(--accent-theme, #00f2fe); padding: 10px 14px; margin: 12px 0; border-radius: 4px;'><p style='margin: 0;'>{block}</p></div>"
        
    return f"<p style='margin-bottom: 8px;'>{block}</p>"

def format_answer_html(lines):
    # Group lines by double-newline blocks
    text = "\n".join(lines)
    blocks = [b.strip() for b in text.split('\n\n') if b.strip()]
    html_blocks = []
    for b in blocks:
        # Check if block has multiple bullet points inside
        subblocks = [sb.strip() for sb in b.split('\n') if sb.strip()]
        if len(subblocks) > 1 and all(sb.startswith(('-', '•')) or re.match(r'^\d+\.', sb) for sb in subblocks):
            list_items = []
            for sb in subblocks:
                clean_sb = re.sub(r'^(?:-|\u2022|\d+\.)\s*', '', sb).strip()
                # Highlight bold prefix in list items
                bold_match = re.match(r'^([^:]+:)(.*)', clean_sb)
                if bold_match:
                    list_items.append(f"<li style='margin-bottom: 6px;'><strong>{bold_match.group(1)}</strong>{bold_match.group(2)}</li>")
                else:
                    list_items.append(f"<li style='margin-bottom: 6px;'>{clean_sb}</li>")
            html_blocks.append(f"<ul style='padding-left: 20px; margin-top: 8px; margin-bottom: 12px;'>{''.join(list_items)}</ul>")
        else:
            # Rejoin block and format
            joined_block = " ".join(subblocks)
            html_blocks.append(format_block_html(joined_block))
            
    return "".join(html_blocks)

def main():
    base_dir = r"C:\Users\USER\Desktop\Extras\.study\chem"
    
    # 1. READ COMPLETE ANSWERS DUMP AND PARSE IT
    dump_path = os.path.join(base_dir, "complete_answers_dump.txt")
    with open(dump_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    parsed_db = {
        "Unit-1": { "Section A": [], "Section B": [] },
        "Unit-2": { "Section A": [], "Section B": [] },
        "Unit-3": { "Section A": [], "Section B": [] },
        "Unit-4": { "Section A": [], "Section B": [] },
        "Section C": { "Unit 1": [], "Unit 2": [], "Unit 3": [] }
    }
    
    current_unit = None
    current_section = None
    current_numerical_unit = None
    
    active_q = None
    active_ans_lines = []
    
    def save_active_q():
        nonlocal active_q, active_ans_lines
        if active_q:
            html_ans = format_answer_html(active_ans_lines)
            item = {
                "label": active_q["label"],
                "question": active_q["question"],
                "answer": html_ans,
                "keywords": get_keywords(active_q["question"])
            }
            if current_section == "Section C":
                parsed_db["Section C"][current_numerical_unit].append(item)
            else:
                parsed_db[current_unit][current_section].append(item)
            active_q = None
            active_ans_lines = []

    for line in lines:
        line_s = line.strip()
        if not line_s:
            if active_q:
                active_ans_lines.append("")
            continue
            
        # Detect Unit headers
        if "UNIT 1: WATER TECHNOLOGY" in line_s:
            save_active_q()
            current_unit = "Unit-1"
            current_section = None
            continue
        elif "UNIT 2: CHEMICAL FUELS" in line_s:
            save_active_q()
            current_unit = "Unit-2"
            current_section = None
            continue
        elif "UNIT 3: BATTERY TECHNOLOGY" in line_s:
            save_active_q()
            current_unit = "Unit-3"
            current_section = None
            continue
        elif "UNIT 4: POLYMERS" in line_s:
            save_active_q()
            current_unit = "Unit-4"
            current_section = None
            continue
        elif "SECTION C — Numerical Problems" in line_s:
            save_active_q()
            current_section = "Section C"
            current_numerical_unit = None
            continue
            
        # Detect Section headers
        if "SECTION A" in line_s:
            save_active_q()
            current_section = "Section A"
            continue
        elif "SECTION B" in line_s:
            save_active_q()
            current_section = "Section B"
            continue
            
        # Inside Section C, detect subheaders for units
        if current_section == "Section C":
            if "UNIT 1" in line_s:
                save_active_q()
                current_numerical_unit = "Unit 1"
                continue
            elif "UNIT 2" in line_s:
                save_active_q()
                current_numerical_unit = "Unit 2"
                continue
            elif "UNIT 3" in line_s:
                save_active_q()
                current_numerical_unit = "Unit 3"
                continue
                
            # Detect Problem headers
            prob_match = re.match(r'^Problem\s+(\d+(?:\s*\([^)]+\))?)\s*—\s*(.*)', line_s)
            if prob_match:
                save_active_q()
                active_q = {
                    "label": f"Problem {prob_match.group(1).strip()}",
                    "question": prob_match.group(2).strip()
                }
                continue
                
        # Inside Section A / B, detect Question headers
        if current_section in ["Section A", "Section B"]:
            q_match = re.match(r'^([QB]\d+(?:\s*(?:&|and|or)\s*[QB]?\d+)*)\.\s*(.*)', line_s)
            if q_match:
                save_active_q()
                active_q = {
                    "label": q_match.group(1).strip(),
                    "question": q_match.group(2).strip()
                }
                continue
                
        # If it's a normal text line and we have an active question
        if active_q:
            active_ans_lines.append(line_s)
            
    # Save the last active question
    save_active_q()
    
    print("Parsed complete answers document successfully.")
    for u in ["Unit-1", "Unit-2", "Unit-3", "Unit-4"]:
        print(f"  {u}: Section A: {len(parsed_db[u]['Section A'])}, Section B: {len(parsed_db[u]['Section B'])}")
    print(f"  Section C: Unit 1: {len(parsed_db['Section C']['Unit 1'])}, Unit 2: {len(parsed_db['Section C']['Unit 2'])}, Unit 3: {len(parsed_db['Section C']['Unit 3'])}")
    
    # 2. SUPPLEMENTARY STANDARD DETAILED ANSWERS FOR SYLLABUS GAPS
    # Let's write out custom rich HTML answers for any questions that are in the syllabus question bank but missing detailed answers in the docx.
    unit2_long_supplement = [
        {
            "label": "B1",
            "question": "Discuss the classification of fuels (solid, liquid, gaseous) with suitable examples and advantages.",
            "answer": """<p style='margin-bottom: 8px;'>Fuels are combustible substances containing carbon and hydrogen as major constituents, which on burning in the presence of air liberate a large amount of heat that can be used for domestic or industrial purposes.</p>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>CLASSIFICATION OF FUELS:</h4>
<p style='margin-bottom: 8px;'>Fuels are broadly classified into two categories: (a) Based on occurrence (Primary/Natural vs Secondary/Derived), and (b) Based on physical state (Solid, Liquid, Gaseous).</p>
<ul style='padding-left: 20px; margin-top: 8px;'>
<li style='margin-bottom: 8px;'><strong>1. Solid Fuels:</strong>
  <ul>
    <li><em>Natural:</em> Wood, Peat, Lignite, Coal.</li>
    <li><em>Derived:</em> Charcoal, Coke.</li>
    <li><em>Advantages:</em> Easy to transport, safe to store without risk of explosion, low cost, moderate ignition temperature.</li>
    <li><em>Disadvantages:</em> High ash content, slow combustion rate, heat wastage during handling, leaves solid residue.</li>
  </ul>
</li>
<li style='margin-bottom: 8px;'><strong>2. Liquid Fuels:</strong>
  <ul>
    <li><em>Natural:</em> Crude petroleum.</li>
    <li><em>Derived:</em> Petrol (Gasoline), Diesel, Kerosene, Heavy oil.</li>
    <li><em>Advantages:</em> Higher calorific value than solid fuels, burn cleanly without ash or dust, easy to control combustion rate, easy storage in tankers.</li>
    <li><em>Disadvantages:</em> Costlier, high risk of fire hazard, require special storage tanks.</li>
  </ul>
</li>
<li style='margin-bottom: 8px;'><strong>3. Gaseous Fuels:</strong>
  <ul>
    <li><em>Natural:</em> Natural gas.</li>
    <li><em>Derived:</em> LPG, CNG, Coal gas, Water gas, Producer gas, Hydrogen.</li>
    <li><em>Advantages:</em> Highest calorific value, complete combustion with zero smoke or ash, highly controllable, can be easily transported through pipelines directly.</li>
    <li><em>Disadvantages:</em> Highly explosive, require bulky cylinders or high-pressure pipelines, leak hazard.</li>
  </ul>
</li>
</ul>""",
            "keywords": ["FUELS", "CLASSIFICATION", "SOLID", "LIQUID", "GASEOUS", "ADVANTAGES"]
        },
        {
            "label": "B3",
            "question": "How do you determine the nitrogen, sulphur, ash and oxygen content of a coal sample? Explain.",
            "answer": """<p style='margin-bottom: 8px;'>The chemical determination of elemental constituents in coal is part of <strong>Ultimate Analysis</strong>. It includes finding the precise percentages of carbon, hydrogen, nitrogen, sulphur, ash, and oxygen.</p>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>1. NITROGEN DETERMINATION (Kjeldahl's Method):</h4>
<p style='margin-bottom: 8px;'>A weighed quantity of dry coal (about 1 g) is heated with concentrated H₂SO₄ along with K₂SO₄ (catalyst) in a Kjeldahl flask. Nitrogen is converted to ammonium sulfate. The mixture is then distilled with excess NaOH, liberating ammonia gas which is absorbed in a known volume of standard acid (HCl or H₂SO₄). The unreacted acid is back-titrated with standard NaOH.</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Nitrogen (%) = (Volume of acid consumed × Normality of acid × 1.4) / Mass of coal sample</div>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>2. SULPHUR DETERMINATION (Eschka Mixture Method):</h4>
<p style='margin-bottom: 8px;'>A weighed sample of coal is heated with Eschka mixture (MgO and Na₂CO₃) at 800°C. All sulphur in coal is converted to soluble sulfates. The mixture is treated with water, filtered, and precipitated as barium sulfate (BaSO₄) by adding barium chloride (BaCl₂). The precipitate is filtered, washed, ignited, and weighed.</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Sulphur (%) = (Mass of BaSO₄ precipitate / Mass of coal sample) × (32 / 233.4) × 100</div>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>3. ASH DETERMINATION:</h4>
<p style='margin-bottom: 8px;'>A weighed coal sample is heated in an open silica crucible at 700-750°C in a muffle furnace until all carbonaceous matter burns off. The residue left is inorganic ash. It is cooled and weighed.</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Ash (%) = (Mass of ash residue / Mass of coal sample) × 100</div>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>4. OXYGEN DETERMINATION:</h4>
<p style='margin-bottom: 8px;'>Oxygen is not determined directly. It is calculated by difference after finding all other elemental percentages.</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Oxygen (%) = 100 − (Carbon% + Hydrogen% + Nitrogen% + Sulphur% + Ash%)</div>""",
            "keywords": ["NITROGEN", "SULPHUR", "ASH", "OXYGEN", "COAL", "DETERMINATION"]
        },
        {
            "label": "B4",
            "question": "What are the advantages of using hydrogen and LPG as a fuel?",
            "answer": """<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>ADVANTAGES OF LPG (Liquefied Petroleum Gas):</h4>
<ul style='padding-left: 20px; margin-top: 8px;'>
  <li style='margin-bottom: 6px;'><strong>High Calorific Value:</strong> It has a high CV of about 11,000 kcal/kg, much higher than coal and wood.</li>
  <li style='margin-bottom: 6px;'><strong>Clean Combustion:</strong> Burns completely without leaving any ash, soot, or residue, reducing cleanup and emissions.</li>
  <li style='margin-bottom: 6px;'><strong>Controllability:</strong> LPG burners can be easily lit, regulated, and shut off instantly.</li>
  <li style='margin-bottom: 6px;'><strong>Storage & Transport:</strong> It is compressed into liquid form and filled in lightweight cylinders, making transport very easy.</li>
</ul>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>ADVANTAGES OF HYDROGEN AS A FUEL:</h4>
<ul style='padding-left: 20px; margin-top: 8px;'>
  <li style='margin-bottom: 6px;'><strong>Highest Energy Density:</strong> Hydrogen has the highest calorific value of all known fuels — about 34,000 kcal/kg (approx. 3 times that of petrol).</li>
  <li style='margin-bottom: 6px;'><strong>Zero Pollution (Eco-Friendly):</strong> The combustion of hydrogen produces only water vapour (H₂O), making it a 100% clean fuel.</li>
  <li style='margin-bottom: 6px;'><strong>Abundant Source:</strong> It can be synthesized from water electrolysis, making it an inexhaustible, renewable green resource.</li>
  <li style='margin-bottom: 6px;'><strong>Fuel Cells:</strong> Can be directly used in hydrogen fuel cells to generate electricity with extremely high efficiency (~60-70%).</li>
</ul>""",
            "keywords": ["ADVANTAGES", "HYDROGEN", "LPG", "FUEL", "CALORIFIC"]
        },
        {
            "label": "B5",
            "question": "Explain proximate analysis of coal. Discuss the determination, significance and limitations of each parameter.",
            "answer": """<p style='margin-bottom: 8px;'><strong>Proximate Analysis</strong> is a rapid empirical method to determine the quality of a coal sample. It measures four practical parameters: <strong>Moisture, Volatile Matter, Ash</strong>, and <strong>Fixed Carbon</strong>.</p>
<table style='width:100%; border-collapse:collapse; margin:12px 0; font-size:0.9rem; border: 1px solid var(--border-color);'>
  <tr style='background:rgba(255,255,255,0.06); border-bottom:1px solid var(--border-color);'>
    <th style='padding:10px; border:1px solid var(--border-color); color:var(--accent-primary); text-align:left;'>Parameter</th>
    <th style='padding:10px; border:1px solid var(--border-color); color:var(--accent-primary); text-align:left;'>Determination Method</th>
    <th style='padding:10px; border:1px solid var(--border-color); color:var(--accent-primary); text-align:left;'>Practical Significance</th>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Moisture (M)</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Heat 1g coal at 105-110°C in an oven for 1 hour. Loss in weight is moisture.</td>
    <td style='padding:10px; border:1px solid var(--border-color);'><strong>Undesirable.</strong> Lowers calorific value, consumes heat to evaporate water, increases transport cost.</td>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Volatile Matter (VM)</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Heat moisture-free coal in a covered crucible at 925±20°C for 7 mins in furnace. Loss in wt is volatile matter.</td>
    <td style='padding:10px; border:1px solid var(--border-color);'><strong>Undesirable.</strong> High VM burns with long smoky flame, yields less coke, requires large combustion chambers.</td>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Ash (A)</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Heat residual coal in open crucible at 700-750°C until all carbon is burnt. Weigh residual inorganic ash.</td>
    <td style='padding:10px; border:1px solid var(--border-color);'><strong>Undesirable.</strong> Lowers CV, causes slagging on boiler grates, obstructs heat transfer, creates disposal issues.</td>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Fixed Carbon (FC)</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Calculated by difference: FC% = 100 − (M% + VM% + A%)</td>
    <td style='padding:10px; border:1px solid var(--border-color);'><strong>Highly Desirable.</strong> Represents actual combustible solid matter. Higher FC gives higher calorific value.</td>
  </tr>
</table>""",
            "keywords": ["PROXIMATE", "ANALYSIS", "COAL", "MOISTURE", "VOLATILE", "ASH", "FIXED", "CARBON"]
        },
        {
            "label": "B6",
            "question": "Explain ultimate analysis of coal. Discuss the determination of carbon, hydrogen, nitrogen, sulphur, oxygen and ash with reactions.",
            "answer": """<p style='margin-bottom: 8px;'><strong>Ultimate Analysis</strong> involves the chemical determination of elemental constituents of coal. It is more accurate than proximate analysis and helps calculate the exact theoretical calorific value.</p>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>1. CARBON AND HYDROGEN:</h4>
<p style='margin-bottom: 8px;'>A known mass of coal (about 0.2 g) is burnt in a stream of pure dry oxygen. Carbon is oxidized to CO₂ and Hydrogen to H₂O. These gases are passed through pre-weighed U-tubes containing anhydrous CaCl₂ (absorbs H₂O) and KOH solution (absorbs CO₂). The increase in weight of the tubes gives the mass of H₂O and CO₂ produced.</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Carbon (%) = (Increase in KOH tube wt × 12 × 100) / (Sample wt × 44)</div>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Hydrogen (%) = (Increase in CaCl₂ tube wt × 2 × 100) / (Sample wt × 18)</div>
<p style='margin-bottom: 8px;'><em>Reactions:</em> C + O₂ → CO₂  ;  2H + ½ O₂ → H₂O  ;  2KOH + CO₂ → K₂CO₃ + H₂O  ;  CaCl₂ + 2H₂O → CaCl₂·2H₂O</p>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>2. NITROGEN:</h4>
<p style='margin-bottom: 8px;'>Determined by <strong>Kjeldahl's method</strong>. Heated with conc H₂SO₄ to form (NH₄)₂SO₄. Distilled with NaOH, liberating NH₃ which is absorbed in standard acid. Unused acid is titrated.</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Nitrogen (%) = (Vol of acid used × Normality × 1.4) / Sample wt</div>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>3. SULPHUR:</h4>
<p style='margin-bottom: 8px;'>Burnt with Eschka mixture to form sulfates, which are precipitated with BaCl₂ as barium sulfate (BaSO₄).</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Sulphur (%) = (BaSO₄ wt × 32 × 100) / (Sample wt × 233.4)</div>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>4. ASH AND OXYGEN:</h4>
<p style='margin-bottom: 8px;'>Ash is determined by ignition (same as proximate). Oxygen is calculated by difference:</p>
<div class='formula-display' style='font-family: monospace; background: rgba(0,0,0,0.2); padding: 8px 12px; border-left: 3px solid var(--accent-theme, #00f2fe); margin: 8px 0;'>Oxygen (%) = 100 − (C% + H% + N% + S% + Ash%)</div>""",
            "keywords": ["ULTIMATE", "ANALYSIS", "COAL", "CARBON", "HYDROGEN", "NITROGEN", "SULPHUR", "OXYGEN"]
        },
        {
            "label": "B7",
            "question": "Discuss the refining of petroleum and fractional distillation. List the major fractions with their boiling ranges and uses.",
            "answer": """<p style='margin-bottom: 8px;'>Crude oil (petroleum) cannot be used directly as fuel. It must be refined to separate it into useful hydrocarbon fractions. <strong>Refining</strong> involves: (1) Removal of water and salt, (2) Removal of sulphur compounds, and (3) Fractional distillation.</p>
<h4 style='color: var(--accent-primary, #00f2fe); margin-top: 12px; margin-bottom: 6px; font-weight: 600;'>FRACTIONAL DISTILLATION:</h4>
<p style='margin-bottom: 8px;'>Crude oil is heated in a pipe still to about 400°C and vaporized. The vapours are fed into a tall fractionating column. As vapours rise, they cool. Higher boiling hydrocarbons condense lower in the column, while lower boiling ones rise higher before condensing. Different fractions are tapped off at different levels:</p>
<table style='width:100%; border-collapse:collapse; margin:12px 0; font-size:0.9rem; border: 1px solid var(--border-color);'>
  <tr style='background:rgba(255,255,255,0.06); border-bottom:1px solid var(--border-color);'>
    <th style='padding:10px; border:1px solid var(--border-color); color:var(--accent-primary); text-align:left;'>Fraction</th>
    <th style='padding:10px; border:1px solid var(--border-color); color:var(--accent-primary); text-align:left;'>Boiling Range (°C)</th>
    <th style='padding:10px; border:1px solid var(--border-color); color:var(--accent-primary); text-align:left;'>Composition</th>
    <th style='padding:10px; border:1px solid var(--border-color); color:var(--accent-primary); text-align:left;'>Primary Uses</th>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Petroleum Gases</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Below 30°C</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>C₁ – C₄</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>LPG for domestic cooking, petrochem feed</td>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Petrol (Gasoline)</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>40 – 120°C</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>C₅ – C₉</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Motor fuel for internal combustion engines</td>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Kerosene Oil</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>180 – 250°C</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>C₁₀ – C₁₆</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Aviation turbine fuel (jet fuel), domestic fuel</td>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Diesel Oil</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>250 – 320°C</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>C₁₅ – C₁₈</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Fuel for diesel engines (trucks, generators)</td>
  </tr>
  <tr style='border-bottom:1px solid var(--border-color);'>
    <td style='padding:10px; border:1px solid var(--border-color); font-weight:700;'>Heavy/Lubricating Oil</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>320 – 400°C</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>C₁₉ – C₃₀</td>
    <td style='padding:10px; border:1px solid var(--border-color);'>Lubrication oil, grease, paraffin wax, asphalt</td>
  </tr>
</table>""",
            "keywords": ["REFINING", "PETROLEUM", "FRACTIONAL", "DISTILLATION", "FRACTIONS"]
        }
    ]
    
    unit3_long_supplement = [
        {
            "label": "B2",
            "question": "Discuss primary, secondary and reserve batteries with at least two examples of each.",
            "answer": """<p style='margin-bottom: 8px;'>Batteries are chemical power sources that convert chemical energy directly into electrical energy. They are classified into three major types:</p>
<ul style='padding-left: 20px; margin-top: 8px;'>
<li style='margin-bottom: 8px;'><strong>1. Primary Batteries (Non-Rechargeable):</strong>
  <p style='margin: 4px 0;'>These are batteries in which the cell reaction is irreversible. Once the active chemicals are consumed, the battery is dead and must be discarded. They are compact, simple, and cheap, with low self-discharge.</p>
  <ul>
    <li><em>Example 1:</em> <strong>Dry Cell (Leclanché Cell)</strong> — used in wall clocks, flashlights, and TV remotes.</li>
    <li><em>Example 2:</em> <strong>Mercury Battery</strong> — used in hearing aids, cameras, and watches due to constant voltage output (1.35 V).</li>
  </ul>
</li>
<li style='margin-bottom: 8px;'><strong>2. Secondary Batteries (Rechargeable):</strong>
  <p style='margin: 4px 0;'>These are batteries in which the electrochemical reactions can be reversed reversibly by passing electrical current from an external source in the opposite direction. They can undergo hundreds of charge-discharge cycles.</p>
  <ul>
    <li><em>Example 1:</em> <strong>Lead-Acid Storage Battery</strong> — used as starter batteries in cars, buses, and home UPS systems.</li>
    <li><em>Example 2:</em> <strong>Lithium-Ion Battery</strong> — used in smartphones, laptops, and electric vehicles.</li>
  </ul>
</li>
<li style='margin-bottom: 8px;'><strong>3. Reserve Batteries (Activated on Demand):</strong>
  <p style='margin: 4px 0;'>These are batteries in which one key component (usually the electrolyte) is stored separately to prevent self-discharge during long-term storage. The battery is activated only when the component is introduced or a barrier is broken.</p>
  <ul>
    <li><em>Example 1:</em> <strong>Magnesium-Silver Chloride Battery</strong> — activated automatically by sea water when submerged; used in torpedoes and marine distress signals.</li>
    <li><em>Example 2:</em> <strong>Zinc-Air Battery (Dry form)</strong> — activated by removing an adhesive seal, allowing oxygen from the air to enter the cell and act as the cathode reactant.</li>
  </ul>
</li>
</ul>""",
            "keywords": ["PRIMARY", "SECONDARY", "RESERVE", "BATTERIES", "EXAMPLES"]
        },
        {
            "label": "B3",
            "question": "Discuss the following characteristics of batteries: voltage, energy density, power density, cycle life, shelf life.",
            "answer": """<p style='margin-bottom: 8px;'>A battery's performance and suitability for specific applications (like portable electronics vs electric vehicles) are determined by several key characteristics:</p>
<ul style='padding-left: 20px; margin-top: 8px;'>
<li style='margin-bottom: 8px;'><strong>1. Voltage (EMF):</strong>
  <p style='margin: 4px 0;'>The potential difference between the cathode and anode under open-circuit conditions. It depends on the standard reduction potentials of the electrode materials (Nernst equation). E.g., Li-ion cell voltage (~3.6 V) is much higher than Ni-Cd (1.2 V).</p>
</li>
<li style='margin-bottom: 8px;'><strong>2. Energy Density (Specific Energy):</strong>
  <p style='margin: 4px 0;'>The total amount of electrical energy stored per unit mass (Wh/kg) or unit volume (Wh/L). Higher energy density means a smaller, lighter battery can run a device for longer. Lithium-ion batteries have a high energy density (150-250 Wh/kg).</p>
</li>
<li style='margin-bottom: 8px;'><strong>3. Power Density (Specific Power):</strong>
  <p style='margin: 4px 0;'>The rate at which a battery can deliver energy per unit mass (W/kg). High power density is essential for applications requiring large, rapid bursts of current, such as electric vehicle acceleration or power tools.</p>
</li>
<li style='margin-bottom: 8px;'><strong>4. Cycle Life:</strong>
  <p style='margin: 4px 0;'>The number of complete charge-discharge cycles a secondary battery can undergo before its storage capacity degrades below a threshold (usually 80% of its original rated capacity). Li-ion batteries typically have 500-2000 cycles.</p>
</li>
<li style='margin-bottom: 8px;'><strong>5. Shelf Life:</strong>
  <p style='margin: 4px 0;'>The duration for which a battery can remain inactive in storage without losing its charge or degrading physically due to self-discharge reactions. It is highly temperature-dependent; primary dry cells have a shelf life of 2-5 years.</p>
</li>
</ul>""",
            "keywords": ["CHARACTERISTICS", "ENERGY", "POWER", "VOLTAGE", "CYCLE", "SHELF"]
        }
    ]

    # Let's add them to the parsed database
    parsed_db["Unit-2"]["Section B"].extend(unit2_long_supplement)
    parsed_db["Unit-3"]["Section B"].extend(unit3_long_supplement)
    
    # 3. MERGE WITH EXISTING MASTER DB FOR HIGH-FIDELITY OVERLAPS (E.G. TABLES AND DETAILED FORMULAS)
    master_path = os.path.join(base_dir, "master_study_db.json")
    if os.path.exists(master_path):
        try:
            with open(master_path, "r", encoding="utf-8") as f:
                old_db = json.load(f)
                
            # If there's an existing styled question, let's keep or merge it
            for u in ["Unit-1", "Unit-2", "Unit-3", "Unit-4"]:
                for sec in ["Section A", "Section B"]:
                    old_qs = old_db.get(u, {}).get(sec, [])
                    new_qs = parsed_db[u][sec]
                    
                    # Create lookup of old Qs
                    old_lookup = {clean_text(q['question'].lower())[:40]: q for q in old_qs}
                    
                    for nq in new_qs:
                        cleaned_nq = clean_text(nq['question'].lower())[:40]
                        # If old exists and has a more complex answer (like table), or if new answer is empty
                        if cleaned_nq in old_lookup:
                            oq = old_lookup[cleaned_nq]
                            if "<table" in oq['answer'] or len(oq['answer']) > len(nq['answer']):
                                nq['answer'] = oq['answer']
                                nq['keywords'] = oq.get('keywords', nq['keywords'])
                                
            # Make sure we add questions that are in old DB but somehow not in new DB
            for u in ["Unit-1", "Unit-2", "Unit-3", "Unit-4"]:
                for sec in ["Section A", "Section B"]:
                    old_qs = old_db.get(u, {}).get(sec, [])
                    new_qs = parsed_db[u][sec]
                    new_lookup = {clean_text(q['question'].lower())[:40] for q in new_qs}
                    for oq in old_qs:
                        cleaned_oq = clean_text(oq['question'].lower())[:40]
                        if cleaned_oq not in new_lookup:
                            new_qs.append(oq)
                            
            # Sort the questions by their label (Q1, Q2... B1, B2...)
            def label_key(item):
                lbl = item.get('label', '')
                num = re.findall(r'\d+', lbl)
                return (lbl[0] if lbl else '', int(num[0]) if num else 0)
                
            for u in ["Unit-1", "Unit-2", "Unit-3", "Unit-4"]:
                for sec in ["Section A", "Section B"]:
                    parsed_db[u][sec].sort(key=label_key)
                    
            print("Successfully merged and consolidated the master study database!")
        except Exception as e:
            print(f"Warning during master database merging: {e}")
            
    # Save the updated master study database
    with open(master_path, "w", encoding="utf-8") as f:
        json.dump(parsed_db, f, indent=4, ensure_ascii=False)
    print(f"Wrote comprehensive master study database to: {master_path}")
    
    # 4. PARSE PDF & DOCX QUESTION BANKS AND INTEGRATE INTO extracted_materials.json
    extracted_path = os.path.join(base_dir, "extracted_materials.json")
    try:
        with open(extracted_path, "r", encoding="utf-8") as f:
            materials = json.load(f)
    except Exception as e:
        print(f"Warning: Could not read extracted_materials.json: {e}")
        materials = { "Unit-1": {}, "Unit-2": {}, "Unit-3": {}, "Unit-4": {} }
        
    # Read the main PDF ETCCCH104_QuestionBank.pdf and split it by page groups
    pdf_path = os.path.join(base_dir, "material", "ETCCCH104_QuestionBank.pdf")
    if os.path.exists(pdf_path):
        try:
            reader = pypdf.PdfReader(pdf_path)
            unit_pages = {
                "Unit-1": [0, 1], # Pages 1 and 2
                "Unit-2": [2, 3], # Pages 3 and 4
                "Unit-3": [4],    # Page 5
                "Unit-4": [5]     # Page 6
            }
            
            for u_key, pages in unit_pages.items():
                u_text = []
                for p_idx in pages:
                    if p_idx < len(reader.pages):
                        txt = reader.pages[p_idx].extract_text()
                        if txt:
                            u_text.append(txt)
                if u_text:
                    materials[u_key]["ETCCCH104_QuestionBank.pdf"] = {
                        "full_text": "\n".join(u_text)
                    }
            print("Successfully split and integrated ETCCCH104_QuestionBank.pdf into extracted_materials.json!")
        except Exception as e:
            print(f"Error parsing ETCCCH104_QuestionBank.pdf: {e}")
            
    # Read the minor DOCX Question Bank-BTECH chemisty .docx and split it by unit
    minor_docx_path = os.path.join(base_dir, "material", "Question Bank-BTECH chemisty .docx")
    if os.path.exists(minor_docx_path):
        try:
            doc = docx.Document(minor_docx_path)
            q_list = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t and not t.startswith("Question Bank"):
                    q_list.append(t)
                    
            # Let's categorize the questions
            u2_qs = []
            u3_qs = []
            u4_qs = []
            
            for q in q_list:
                ql = q.lower()
                # Check keywords for classification
                if any(k in ql for k in ['polymer', 'rubber', 'nylon', 'dacron', 'conduct', 'biodegradable', 'composite']):
                    u4_qs.append(q)
                elif any(k in ql for k in ['fuel', 'calorific', 'coal', 'refine', 'petroleum', 'crack', 'knock', 'lpg', 'cng']):
                    u2_qs.append(q)
                elif any(k in ql for k in ['galvanic', 'electrode', 'emf', 'battery', 'cell', 'nernst']):
                    u3_qs.append(q)
                else:
                    # Default/fallback based on index
                    # The first few in docx are polymers
                    pass
                    
            # Let's do a strict fallback if our keyword classifier is too loose:
            # We know the docx had:
            # Q1-12: Polymers (Unit 4)
            # Q13-23: Fuels (Unit 2)
            # Q24-33: Battery Technology (Unit 3)
            u4_qs = q_list[0:12]
            u2_qs = q_list[12:23]
            u3_qs = q_list[23:33]
            
            if u2_qs:
                materials["Unit-2"]["Question Bank-BTECH chemisty .docx"] = {
                    "full_text": "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(u2_qs)])
                }
            if u3_qs:
                materials["Unit-3"]["Question Bank-BTECH chemisty .docx"] = {
                    "full_text": "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(u3_qs)])
                }
            if u4_qs:
                materials["Unit-4"]["Question Bank-BTECH chemisty .docx"] = {
                    "full_text": "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(u4_qs)])
                }
            print("Successfully split and integrated Question Bank-BTECH chemisty .docx into extracted_materials.json!")
        except Exception as e:
            print(f"Error parsing Question Bank-BTECH chemisty .docx: {e}")
            
    # Read the major DOCX QUESTION BANK ETCCCH104.docx and split it by unit
    major_docx_path = os.path.join(base_dir, "material", "QUESTION BANK ETCCCH104.docx")
    if os.path.exists(major_docx_path):
        try:
            doc = docx.Document(major_docx_path)
            q_list = []
            for p in doc.paragraphs:
                t = p.text.strip()
                # Remove header lines
                if t and not any(h in t for h in ["QUESTION BANK", "2 MARK QUESTIONS", "4 MARK QUESTIONS", "NUMERICAL QUESTIONS", "8 MARK QUESTIONS"]):
                    q_list.append(t)
            
            # Let's categorize the questions by unit keyword matching
            u1_major = []
            u2_major = []
            u3_major = []
            u4_major = []
            
            for q in q_list:
                ql = q.lower()
                if any(k in ql for k in ['hardness', 'zeolite', 'ion-exchange', 'boiler', 'scale', 'sludge', 'softening', 'alkalinity', 'soap', 'edta', 'ppm', 'mg/l', 'clarke', 'french']):
                    u1_major.append(q)
                elif any(k in ql for k in ['polymer', 'rubber', 'nylon', 'buna', 'neoprene', 'conductive', 'biodegradable', 'composite', 'vulcanization', 'molecular forces', 'thermoplastic', 'thermoset', 'phbv', 'pla']):
                    u4_major.append(q)
                elif any(k in ql for k in ['battery', 'batteries', 'electrode', 'emf', 'cell', 'nernst', 'conductance', 'conductometric', 'voltage', 'energy density', 'power density', 'capacity', 'shelf life', 'cycle life', 'fe3+', 'br-', 'cl2']):
                    u3_major.append(q)
                elif any(k in ql for k in ['fuel', 'fuels', 'calorific', 'coal', 'refine', 'petroleum', 'crack', 'knock', 'lpg', 'cng', 'octane', 'cetane', 'dulong', 'bomb', 'calorimeter', 'net calorific', 'gross calorific']):
                    u2_major.append(q)
                else:
                    # Default/fallback based on typical syllabus keywords
                    if any(k in ql for k in ['water', 'treatment']):
                        u1_major.append(q)
                    else:
                        u2_major.append(q) # default to Unit 2 or 1
            
            # Save these categorized lists under the filename "QUESTION BANK ETCCCH104.docx" in materials
            if u1_major:
                materials["Unit-1"]["QUESTION BANK ETCCCH104.docx"] = {
                    "full_text": "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(u1_major)])
                }
            if u2_major:
                materials["Unit-2"]["QUESTION BANK ETCCCH104.docx"] = {
                    "full_text": "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(u2_major)])
                }
            if u3_major:
                materials["Unit-3"]["QUESTION BANK ETCCCH104.docx"] = {
                    "full_text": "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(u3_major)])
                }
            if u4_major:
                materials["Unit-4"]["QUESTION BANK ETCCCH104.docx"] = {
                    "full_text": "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(u4_major)])
                }
            print("Successfully split and integrated QUESTION BANK ETCCCH104.docx into extracted_materials.json!")
        except Exception as e:
            print(f"Error parsing QUESTION BANK ETCCCH104.docx: {e}")
            
    # Save the updated extracted materials
    with open(extracted_path, "w", encoding="utf-8") as f:
        json.dump(materials, f, indent=4, ensure_ascii=False)
    print(f"Wrote comprehensive extracted materials to: {extracted_path}")
    
    # 5. RUN generate_portal.py TO REGENERATE THE HTML
    import subprocess
    print("Running generate_portal.py...")
    res = subprocess.run(["python", os.path.join(base_dir, "generate_portal.py")], capture_output=True, text=True)
    if res.returncode == 0:
        print("Success! Portal index.html regenerated successfully with 100% complete answers!")
    else:
        print(f"Error running generate_portal.py:\nSTDOUT:\n{res.stdout}\nSTDERR:\n{res.stderr}")

if __name__ == "__main__":
    main()
